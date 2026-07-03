"""File loaders for PDF, DOCX, MD/TXT into LangChain Documents.

Each loader returns a list of :class:`langchain_core.documents.Document`
with metadata that downstream pieces (chunker, vectorstore, retriever) rely on.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable, List

from langchain_core.documents import Document

from src.utils.logger import get_logger

log = get_logger(__name__)


SUPPORTED_EXT = {".pdf", ".docx", ".md", ".markdown", ".txt"}


# --------------------------------------------------------------------------- #
# Individual format loaders
# --------------------------------------------------------------------------- #
def _load_pdf(path: Path) -> List[Document]:
    """Load a PDF. Uses pypdf so we keep dependencies light."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    docs: List[Document] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - very rare
            log.warning("Failed to extract text from %s page %d: %s", path, i, exc)
            text = ""
        text = text.strip()
        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "page": i + 1,
                        "ext": ".pdf",
                    },
                )
            )
    return docs


def _load_docx(path: Path) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    # also pull table cells (resumes often use tables for layout)
    table_parts: list[str] = []
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                table_parts.append(row_text)
    if table_parts:
        text += "\n\n" + "\n".join(table_parts)

    text = text.strip()
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": str(path), "filename": path.name, "ext": ".docx"},
        )
    ]


def _load_text(path: Path) -> List[Document]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={"source": str(path), "filename": path.name, "ext": path.suffix.lower()},
        )
    ]


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def load_document(path: Path | str, category: str | None = None) -> List[Document]:
    """Dispatch to the correct loader based on file extension.

    Args:
        path: file path on disk.
        category: optional tag stored on each document's metadata
            (e.g. "resume", "thesis", "paper"). Useful for filtered retrieval.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    ext = path.suffix.lower()
    if ext == ".pdf":
        docs = _load_pdf(path)
    elif ext == ".docx":
        docs = _load_docx(path)
    elif ext in {".md", ".markdown", ".txt"}:
        docs = _load_text(path)
    else:
        log.warning("Skipping unsupported file: %s", path)
        return []

    if category:
        for d in docs:
            d.metadata["category"] = category

    log.info("Loaded %d block(s) from %s", len(docs), path.name)
    return docs


def load_directory(
    directory: Path | str,
    category: str | None = None,
    recursive: bool = True,
) -> List[Document]:
    """Load every supported file under `directory`."""
    directory = Path(directory)
    if not directory.exists():
        log.warning("Directory does not exist: %s", directory)
        return []

    iterator: Iterable[Path] = (
        directory.rglob("*") if recursive else directory.glob("*")
    )
    out: List[Document] = []
    for p in iterator:
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXT:
            out.extend(load_document(p, category=category))
    return out
